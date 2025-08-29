import os
import csv
import json
import hashlib
from typing import Dict, List, Optional
from dotenv import load_dotenv

from langchain.schema import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader
from docx import Document as DocxDocument

# ---- Embeddings (Hugging Face, local) ----
from langchain_huggingface.embeddings import HuggingFaceEmbeddings

# ---- Vector store (Chroma) - use the new integration ONLY ----
from langchain_chroma import Chroma
from langchain_community.vectorstores.chroma import Chroma


# ---- LLM (Google Gemini) ----
from langchain_google_genai import ChatGoogleGenerativeAI

# ---- Retrieval chain (LangChain 0.3) ----
from langchain_core.prompts import PromptTemplate
from langchain.chains import create_retrieval_chain
from langchain.chains.combine_documents import create_stuff_documents_chain

load_dotenv()
# quiet Chroma telemetry
os.environ.setdefault("ANONYMIZED_TELEMETRY", "False")

# --- Config ---
CHROMA_DIR = os.getenv("CHROMA_DIR", "./chroma_db")
COLLECTION = os.getenv("COLLECTION", "sjsu_campus_copilot")

# You wanted to keep the key inline for now:
os.environ["GOOGLE_API_KEY"] = os.environ.get(
    "GOOGLE_API_KEY", "AIzaSyA5ApL2ltwbFn4GQShKTz_t35744kkq8dg")
GOOGLE_KEY = os.environ["GOOGLE_API_KEY"]
LLM_MODEL = os.getenv("LLM_MODEL", "gemini-1.5-flash")

# ---- Simple manifest to avoid re-indexing unchanged files ----
MANIFEST_PATH = os.path.join(CHROMA_DIR, "manifest.json")


def _load_manifest() -> Dict[str, Dict]:
    try:
        with open(MANIFEST_PATH, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {}


def _save_manifest(m: Dict[str, Dict]):
    os.makedirs(CHROMA_DIR, exist_ok=True)
    tmp = MANIFEST_PATH + ".tmp"
    with open(tmp, "w", encoding="utf-8") as f:
        json.dump(m, f)
    os.replace(tmp, MANIFEST_PATH)


def _file_sha1(path: str) -> str:
    h = hashlib.sha1()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


# --- Embeddings (MiniLM, 384-dim, fast & solid) ---
embeddings = HuggingFaceEmbeddings(
    model_name="sentence-transformers/all-MiniLM-L6-v2")

# --- Vector store (persisted locally) ---
vs = Chroma(
    collection_name=COLLECTION,
    embedding_function=embeddings,
    persist_directory=CHROMA_DIR,
)

# --- LLM (Gemini chat) ---
llm = ChatGoogleGenerativeAI(
    model=LLM_MODEL,
    api_key=GOOGLE_KEY,          # IMPORTANT: use api_key= + REST transport to avoid ADC
    transport="rest",
    temperature=0.2,
    max_output_tokens=512,
    convert_system_message_to_human=True,
    timeout=60,
)

# --- RAG chain ---
PROMPT = PromptTemplate.from_template(
    "You are a helpful assistant. Answer ONLY from the context. "
    "If the answer is not in the context, say you don't know.\n\n"
    "Context:\n{context}\n\nQuestion: {input}\nAnswer:"
)
docs_chain = create_stuff_documents_chain(llm, PROMPT)
# retriever = vs.as_retriever(search_type="mmr", search_kwargs={
#                             "k": 6, "fetch_k": 24, "lambda_mult": 0.2})
# rag_chain = create_retrieval_chain(retriever, docs_chain)
retriever = vs.as_retriever(
    search_type="mmr",
    search_kwargs={"k": 6, "fetch_k": 32, "lambda_mult": 0.2}
)
rag_chain = create_retrieval_chain(retriever, docs_chain)


# ---------- Loaders ----------
SPLITTER = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=150)
TEXT_EXTS = {".txt", ".md", ".markdown", ".log"}
CSV_EXTS = {".csv"}
ACCEPTED_EXTS = {".pdf", ".docx", ".txt", ".md", ".markdown", ".log", ".csv"}


# def load_pdf_to_docs(path: str) -> List[Document]:
#     docs: List[Document] = []
#     reader = PdfReader(path)
#     for i, p in enumerate(reader.pages, start=1):
#         text = (p.extract_text() or "").strip()
#         if not text:
#             continue
#         for chunk in SPLITTER.split_text(text):
#             docs.append(Document(page_content=chunk, metadata={
#                         "source": os.path.basename(path), "page": i, "ext": ".pdf"}))
#     return docs
try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None


def load_pdf_to_docs(path: str) -> List[Document]:
    docs: List[Document] = []
    reader = PdfReader(path)
    mu = None
    for i, p in enumerate(reader.pages, start=1):
        text = (p.extract_text() or "").strip()
        if not text and fitz is not None:
            mu = mu or fitz.open(path)
            text = (mu.load_page(i - 1).get_text("text") or "").strip()
        if not text:
            continue
        for chunk in SPLITTER.split_text(text):
            docs.append(Document(
                page_content=chunk,
                metadata={"source": os.path.basename(
                    path), "page": i, "ext": ".pdf"}
            ))
    if mu:
        mu.close()
    return docs


def load_docx_to_docs(path: str) -> List[Document]:
    docs: List[Document] = []
    d = DocxDocument(path)
    text = "\n".join([para.text for para in d.paragraphs if para.text])
    if not text.strip():
        return docs
    for chunk in SPLITTER.split_text(text):
        docs.append(Document(page_content=chunk, metadata={
                    "source": os.path.basename(path), "ext": ".docx"}))
    return docs


def load_text_to_docs(path: str) -> List[Document]:
    docs: List[Document] = []
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()
    if not text.strip():
        return docs
    for chunk in SPLITTER.split_text(text):
        docs.append(Document(page_content=chunk, metadata={
                    "source": os.path.basename(path)}))
    return docs


def load_csv_to_docs(path: str) -> List[Document]:
    docs: List[Document] = []
    rows = []
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        for row in csv.reader(f):
            rows.append(", ".join(row))
    text = "\n".join(rows)
    if not text.strip():
        return docs
    for chunk in SPLITTER.split_text(text):
        docs.append(Document(page_content=chunk, metadata={
                    "source": os.path.basename(path), "ext": ".csv"}))
    return docs


def read_file_to_docs(path: str) -> List[Document]:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return load_pdf_to_docs(path)
    if ext == ".docx":
        return load_docx_to_docs(path)
    if ext in TEXT_EXTS:
        return load_text_to_docs(path)
    if ext in CSV_EXTS:
        return load_csv_to_docs(path)
    return []

# ---------- Public API ----------


def ingest_any_file(path: str, source_name: str | None = None) -> int:
    docs = read_file_to_docs(path)
    if not docs:
        return 0

    # Use the real uploaded name in metadata (not tmpxxxxx.pdf)
    if source_name:
        real = os.path.basename(source_name)
        for d in docs:
            d.metadata["source"] = real

    # Replace any older vectors for this file
    try:
        src = docs[0].metadata.get("source") or os.path.basename(path)
        vs.delete(where={"source": src})
    except Exception:
        pass

    vs.add_documents(docs)
    vs.persist()
    return len(docs)


# def ingest_any_file(path: str) -> int:
#     """Idempotent ingest: skip if file unchanged; replace if changed."""
#     src = os.path.basename(path)
#     file_hash = _file_sha1(path)
#     manifest = _load_manifest()
#     existing = manifest.get(src, {})

#     if existing.get("sha1") == file_hash:
#         # unchanged -> skip
#         return 0

#     # Build docs
#     docs = read_file_to_docs(path)
#     if not docs:
#         return 0

#     # ensure metadata has source
#     for d in docs:
#         d.metadata["source"] = src

#     # If file existed before (changed), delete old vectors for this source
#     try:
#         vs.delete(where={"source": src})
#     except Exception:
#         pass

#     # Add and persist
#     vs.add_documents(docs)
#     vs.persist()

#     # Update manifest
#     manifest[src] = {"sha1": file_hash, "count": len(docs)}
#     _save_manifest(manifest)
#     return len(docs)


def add_docs_from_pdf(path: str) -> int:
    # just reuse ingest (it handles idempotency + persist)
    return ingest_any_file(path)


def bootstrap_data_dir(folder: str) -> Dict:
    """Index every accepted file in 'folder' once (idempotent)."""
    total_added = 0
    total_skipped = 0
    if not folder or not os.path.isdir(folder):
        return {"dir": folder, "added": 0, "skipped": 0}

    for root, _, files in os.walk(folder):
        for name in files:
            ext = os.path.splitext(name)[1].lower()
            if ext not in ACCEPTED_EXTS:
                continue
            path = os.path.join(root, name)
            added = ingest_any_file(path)
            if added > 0:
                total_added += added
            else:
                total_skipped += 1

    return {"dir": folder, "added": total_added, "skipped_files": total_skipped}


def list_sources() -> List[str]:
    try:
        col = vs._collection
        out = col.get(include=["metadatas"])
        names = set()
        for m in out.get("metadatas", []) or []:
            if m and m.get("source"):
                names.add(m["source"])
        return sorted(names)
    except Exception:
        return []


def preview_retrieval(query: str, n: int = 5) -> List[Dict]:
    docs = retriever.get_relevant_documents(query)
    out = []
    for d in docs[:n]:
        meta = d.metadata or {}
        out.append({
            "source": meta.get("source"),
            "page": meta.get("page"),
            "snippet": (d.page_content or "")[:300].replace("\n", " ")
        })
    return out


def answer_query(q: str) -> Dict:
    out = rag_chain.invoke({"input": q})

    # Coerce answer to string
    ans = out.get("answer", "")
    answer = getattr(ans, "content", ans)
    answer = str(answer)

    # JSON-safe citations
    ctx = out.get("context") or []
    citations = []
    if isinstance(ctx, list):
        for d in ctx:
            meta = getattr(d, "metadata", {}) or {}
            citations.append(
                {"source": str(meta.get("source") or "doc"), "page": meta.get("page", None)})

    return {"answer": answer, "citations": citations}
